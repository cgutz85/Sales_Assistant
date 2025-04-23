import nltk
import os
import warnings
import tiktoken
from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_ollama import OllamaEmbeddings

from langchain_community.vectorstores import FAISS 
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document

from langfuse import Langfuse
from langfuse.callback import CallbackHandler
from vector_store_prep import number_documents

# Reference for Langfuse Handler: https://langfuse.com/docs/integrations/langchain/tracing
langfuse_handler = CallbackHandler(
  secret_key="sk-lf-0ac561f9-87eb-4777-ba60-3b4962e600fc",
  public_key="pk-lf-a88fa50b-358a-4b31-8545-052a29896066",
  host="http://localhost:3000"
)


embeddings = OllamaEmbeddings(model='mxbai-embed-large', base_url='http://localhost:11434')
# Needs to be adjusted to the actual environment (folder structure) in which the script will be running.
db_name = r"C:\Users\cagut\OneDrive\Desktop\AI Agent - Offer Document Generation\WORD_APP\test_db"
# Reference for vector store loading: https://python.langchain.com/docs/integrations/vectorstores/faiss/ 
vector_store = FAISS.load_local(db_name, embeddings, allow_dangerous_deserialization=True)

question = "I need all the information about the project 'Onboarding Fake Bank'. I need task costs, resources and task lists. For this, please check for task items and their total effort estimates. On top of that I need a detailed description of the project."

docs = vector_store.search(query=question, k=5, search_type="similarity")
#print(docs)

# retriever to get the documents that fulfill the similarity requirements within the vector store
#retriever = vector_store.as_retriever(search_type="similarity_score_threshold", search_kwargs = {'k':3, 'score_threshold': 0.3})

#response = retriever.invoke(question)
#print(response)

# prompt format and defintion for the llm to know what it is supposed to do
prompt = """
You are an eloquent sales manager who needs to write offer documents for potential customers. Please produce your output in German and with a professional tone of voice. Good grammar is important. If you do not find information to answer questions please say so. At the end of your outputs please cite the document where you took the information from.
Context: {context}
Answer: 
"""

# Reference for ChatPromptTemplate: https://python.langchain.com/api_reference/core/prompts/langchain_core.prompts.chat.ChatPromptTemplate.html
prompt = ChatPromptTemplate.from_template(prompt)
#print(prompt)

# llm definition
# Reference for LLM definition: https://python.langchain.com/v0.2/api_reference/community/chat_models/langchain_community.chat_models.ollama.ChatOllama.html 
llm = ChatOllama(model='phi4', base_url='http://localhost:11434', temperature=0.01)
llm.invoke('hi')

# retrived context data is formatted into a string, so it can be used by the llm
# Reference for formatting retrieved data: https://python.langchain.com/v0.1/docs/use_cases/question_answering/sources/
def format_docs(docs):
    return '\n\n'.join([doc.page_content for doc in docs])

# formatted context that will be needed by the llm 
context = format_docs(docs)
#print(context)

# Reference for RAG chain: https://python.langchain.com/v0.2/docs/tutorials/rag/
rag_chain = (
    {"context": retriever|format_docs, "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()
)

question = """Please write a summary of the project 'Onboarding Fake Bank' containing the following paragraphs: 
- A brief summary of the project, 
- The offer price, start date and end date including a task list stating the resource roles and team member names belonging to these roles. Do not provide any other numbers than the offer price
- A brief conclusion thanking the client for the possible future collaboration. 
Your output should not exceed one page.""" 

response = rag_chain.invoke(question, config={"callbacks": [langfuse_handler]})


#print(response)

#   Write the result to the output file
# Reference for docx library: https://python-docx.readthedocs.io/en/latest/
document = Document()
document.add_heading('Offerte Onboarding Fake Bank', level=1)
document.add_paragraph(response)

document.save('./offer_doc/offer.docx')
